"""
Resume Builder Service — Generates ATS-friendly .docx resumes from profile data.

Pulls from UserProfile.extra_data["resume_profile"] (student-editable fields)
and UserProfile base fields (ERP auto-fill) to produce a clean, single-column
Word document optimized for ATS scanners.

Templates:
  - classic: Clean, professional, single-column. Academic style.
"""
import io
import logging
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app import models

logger = logging.getLogger("acadmix.resume_builder")

# ═══════════════════════════════════════════════════════════════════════════════
# Constants
# ═══════════════════════════════════════════════════════════════════════════════
_FONT = "Times New Roman"
_CLR = RGBColor(0x00, 0x00, 0x00)  # Pure black for everything
_CONTENT_WIDTH = Cm(21 - 2.0 - 2.0)  # A4 width minus margins


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _margins(doc: Document):
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)


def _run(p, text, sz=Pt(11), bold=False, italic=False):
    """Add a black run — everything is black in this template."""
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = sz
    r.font.color.rgb = _CLR
    if bold:
        r.bold = True
    if italic:
        r.font.italic = True
    return r


def _para(doc, sp_before=0, sp_after=0, align=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(sp_before)
    fmt.space_after = Pt(sp_after)
    if align:
        p.alignment = align
    return p


def _heading(doc, title):
    """Section heading — bold, larger, with thin bottom border. Not uppercase."""
    p = _para(doc, sp_before=8, sp_after=3)
    _run(p, title, sz=Pt(13), bold=True)

    # Thin bottom border
    pPr = p._p.get_or_add_pPr()
    bdr = pPr.makeelement(qn('w:pBdr'), {})
    bdr.append(bdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    }))
    pPr.append(bdr)
    return p


def _bullet(doc, text):
    """Bullet point — round bullet, indented."""
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    _run(p, text, sz=Pt(11))
    return p


def _bullet_to_cell(cell, text, size=9):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"• {text}")
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = _CLR
    return p


def _cell_heading(cell, title, color="000000"):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def _cell_text(cell, text, size=8, bold=False, italic=False, color="000000"):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.font.italic = italic
    return p


def _tech_stack_line(doc, tech):
    """Tech Stack: label bold, items regular. Indented to match bullets, no bullet marker."""
    p = _para(doc, sp_before=1, sp_after=2)
    p.paragraph_format.left_indent = Cm(1.27)  # Match bullet indent
    _run(p, "Tech Stack: ", sz=Pt(11), bold=True)
    _run(p, tech, sz=Pt(11))
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# Template: Classic
# ═══════════════════════════════════════════════════════════════════════════════

def _build_classic(data: Dict[str, Any]) -> Document:
    doc = Document()
    _margins(doc)

    personal = data.get("personal", {})
    education = data.get("education_history", [])
    skills = data.get("skills", {})
    projects = data.get("projects", [])
    experience = data.get("experience", [])
    certs = data.get("certifications", [])
    achievements = data.get("achievements", [])
    summary = data.get("summary", "")

    # ── NAME ──────────────────────────────────────────────────
    name = personal.get("name", "").strip()
    p = _para(doc, sp_before=0, sp_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    if name:
        _run(p, name, sz=Pt(18), bold=True)
    else:
        _run(p, "[Your Full Name]", sz=Pt(18), bold=True, italic=True)

    # ── CONTACT ───────────────────────────────────────────────
    contact = [v for v in [personal.get("email"), personal.get("phone"), personal.get("location")] if v]
    if contact:
        p = _para(doc, sp_before=0, sp_after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, "  |  ".join(contact), sz=Pt(10))

    # ── LINKS ─────────────────────────────────────────────────
    links = [v for v in [personal.get("linkedin"), personal.get("github"), personal.get("portfolio")] if v]
    if links:
        p = _para(doc, sp_before=0, sp_after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, "  |  ".join(links), sz=Pt(10))

    # ── SUMMARY ───────────────────────────────────────────────
    if summary and summary.strip():
        _heading(doc, "Summary")
        p = _para(doc, sp_before=1, sp_after=2)
        _run(p, summary.strip(), sz=Pt(11))

    # ── EDUCATION ─────────────────────────────────────────────
    _heading(doc, "Education")

    current = personal.get("current_education")
    if current and current.get("institution"):
        branch = current.get("branch", "")
        degree = current.get("degree", "B.Tech")
        title = f"{degree} in {branch}" if branch else degree

        # Line 1: Degree + Institution .... right-aligned Batch
        p = _para(doc, sp_before=2, sp_after=0)
        p.paragraph_format.tab_stops.add_tab_stop(_CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
        _run(p, f"{title}, ", sz=Pt(11), bold=True)
        _run(p, current["institution"], sz=Pt(11))
        if current.get("batch"):
            _run(p, "\t", sz=Pt(11))
            _run(p, current["batch"], sz=Pt(11), italic=True)

    for edu in education:
        degree = edu.get("degree") or edu.get("level", "")
        school = edu.get("school", "")
        location = edu.get("location", "")
        field = edu.get("field") or edu.get("board", "")
        year = edu.get("gradYear") or edu.get("year", "")
        month = edu.get("gradMonth", "")
        percentage = edu.get("percentage", "")

        # Line 1: Degree, School .... right-aligned graduation date
        p = _para(doc, sp_before=4, sp_after=0)
        p.paragraph_format.tab_stops.add_tab_stop(_CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
        _run(p, degree, sz=Pt(11), bold=True)
        if school:
            _run(p, f", {school}", sz=Pt(11))
        grad_date = f"{month} {year}".strip() if month else str(year) if year else ""
        if grad_date:
            _run(p, "\t", sz=Pt(11))
            _run(p, grad_date, sz=Pt(11), italic=True)

        # Line 2: Field of Study | Location | Percentage
        sub = [v for v in [field, location, percentage] if v]
        if sub:
            p2 = _para(doc, sp_before=0, sp_after=1)
            _run(p2, "  |  ".join(sub), sz=Pt(10), italic=True)

    # ── SKILLS ────────────────────────────────────────────────
    has_skills = any(skills.get(k) for k in ["languages", "frameworks", "tools", "databases"])
    if has_skills:
        _heading(doc, "Technical Skills")
        for key, label in [("languages", "Languages"), ("frameworks", "Frameworks"),
                           ("tools", "Tools & Platforms"), ("databases", "Databases")]:
            items = skills.get(key, [])
            if items:
                p = _para(doc, sp_before=0, sp_after=1)
                _run(p, f"{label}: ", sz=Pt(11), bold=True)
                _run(p, ", ".join(items), sz=Pt(11))

    # ── PROJECTS ──────────────────────────────────────────────
    if projects:
        _heading(doc, "Projects")
        for proj in projects:
            title = proj.get("title", "Untitled")

            # Bold project title
            p = _para(doc, sp_before=3, sp_after=1)
            _run(p, title, sz=Pt(11), bold=True)

            # Bullets first
            bullets = proj.get("bullets", [])
            for b in bullets:
                if b and b.strip():
                    _bullet(doc, b.strip())

            # Tech Stack: at the end, indented to match bullets
            tech = proj.get("tech_stack", "")
            if tech:
                _tech_stack_line(doc, tech)

    # ── EXPERIENCE ────────────────────────────────────────────
    if experience:
        _heading(doc, "Experience")
        for exp in experience:
            role = exp.get("role", "Role")
            company = exp.get("company", "")

            # Bold role title
            p = _para(doc, sp_before=3, sp_after=0)
            _run(p, role, sz=Pt(11), bold=True)
            if company:
                _run(p, f", {company}", sz=Pt(11))

            # Duration/location subtitle
            sub = [v for v in [exp.get("duration"), exp.get("location")] if v]
            if sub:
                p2 = _para(doc, sp_before=0, sp_after=1)
                _run(p2, "  |  ".join(sub), sz=Pt(10), italic=True)

            # Bullets
            bullets = exp.get("bullets", [])
            for b in bullets:
                if b and b.strip():
                    _bullet(doc, b.strip())

    # ── CERTIFICATIONS ────────────────────────────────────────
    if certs:
        _heading(doc, "Certifications")
        for c in certs:
            p = _para(doc, sp_before=0, sp_after=1)
            _run(p, c.get("name", ""), sz=Pt(11), bold=True)
            suffix = [v for v in [c.get("issuer"), c.get("year") and str(c["year"])] if v]
            if suffix:
                _run(p, f" — {', '.join(suffix)}", sz=Pt(11))

    # ── ACHIEVEMENTS ──────────────────────────────────────────
    if achievements:
        _heading(doc, "Achievements")
        for a in achievements:
            text = a if isinstance(a, str) else a.get("title", "")
            if text and text.strip():
                _bullet(doc, text.strip())

    return doc


def _build_compact(data: Dict[str, Any]) -> Document:
    doc = Document()
    _margins(doc)
    personal = data.get("personal", {})
    skills = data.get("skills", {})
    projects = data.get("projects", [])
    experience = data.get("experience", [])
    certs = data.get("certifications", [])
    achievements = data.get("achievements", [])

    p = _para(doc, sp_before=0, sp_after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    _run(p, personal.get("name") or "[Your Full Name]", sz=Pt(17), bold=True)
    contact = [personal.get("email"), personal.get("phone"), personal.get("location")]
    _run(_para(doc, sp_after=4), " | ".join([x for x in contact if x]), sz=Pt(9))

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.width = Cm(7.0)
    right.width = Cm(9.5)

    current = personal.get("current_education") or {}
    _cell_heading(left, "Education")
    if current.get("institution"):
        _cell_text(left, f"{current.get('degree', 'B.Tech')} {current.get('branch', '')}".strip(), 8, bold=True)
        _cell_text(left, f"{current.get('institution')} | {current.get('batch', '')}".strip(" |"), 8)
    for edu in data.get("education_history", [])[:3]:
        _cell_text(left, f"{edu.get('degree') or edu.get('level', '')} - {edu.get('school', '')}".strip(" -"), 8, bold=True)

    _cell_heading(left, "Skills")
    for key, label in [("languages", "Languages"), ("frameworks", "Frameworks"), ("tools", "Tools"), ("databases", "Databases")]:
        if skills.get(key):
            _cell_text(left, f"{label}: {', '.join(skills[key])}", 8)

    if certs:
        _cell_heading(left, "Certifications")
        for cert in certs[:5]:
            _cell_text(left, cert.get("name", ""), 8)
    if achievements:
        _cell_heading(left, "Achievements")
        for item in achievements[:5]:
            _bullet_to_cell(left, item if isinstance(item, str) else item.get("title", ""), 8)

    if data.get("summary"):
        _cell_heading(right, "Summary")
        _cell_text(right, data["summary"], 8)
    _cell_heading(right, "Projects")
    for project in projects[:4]:
        _cell_text(right, project.get("title", "Untitled Project"), 9, bold=True)
        for bullet in project.get("bullets", [])[:3]:
            if str(bullet).strip():
                _bullet_to_cell(right, bullet, 8)
        if project.get("tech_stack"):
            _cell_text(right, f"Stack: {project['tech_stack']}", 8, italic=True)
    if experience:
        _cell_heading(right, "Experience")
        for exp in experience[:3]:
            _cell_text(right, f"{exp.get('role', '')}, {exp.get('company', '')}".strip(", "), 9, bold=True)
            for bullet in exp.get("bullets", [])[:3]:
                if str(bullet).strip():
                    _bullet_to_cell(right, bullet, 8)
    return doc


def _build_sidebar(data: Dict[str, Any], template: str) -> Document:
    doc = Document()
    _margins(doc)
    accent = {
        "developer": "2563EB",
        "core-engineering": "0F766E",
        "management": "7C3AED",
        "modern": "111827",
    }.get(template, "111827")
    personal = data.get("personal", {})
    skills = data.get("skills", {})
    projects = data.get("projects", [])
    experience = data.get("experience", [])

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.width = Cm(6.2)
    right.width = Cm(10.4)
    _cell_text(left, personal.get("name") or "[Your Full Name]", 15, bold=True, color=accent)
    _cell_text(left, " | ".join([x for x in [personal.get("email"), personal.get("phone"), personal.get("location")] if x]), 8)
    if data.get("summary"):
        _cell_heading(left, "Profile", accent)
        _cell_text(left, data["summary"], 8)
    _cell_heading(left, "Skills", accent)
    for key, label in [("languages", "Languages"), ("frameworks", "Frameworks"), ("tools", "Tools"), ("databases", "Databases")]:
        if skills.get(key):
            _cell_text(left, f"{label}: {', '.join(skills[key])}", 8)
    current = personal.get("current_education") or {}
    if current.get("institution"):
        _cell_heading(left, "Education", accent)
        _cell_text(left, f"{current.get('degree', 'B.Tech')} {current.get('branch', '')}".strip(), 8, bold=True)
        _cell_text(left, current.get("institution", ""), 8)

    primary_title = "Selected Engineering Projects" if template in {"developer", "core-engineering"} else "Leadership & Experience" if template == "management" else "Experience"
    _cell_heading(right, primary_title, accent)
    if template in {"developer", "core-engineering"}:
        for project in projects[:5]:
            _cell_text(right, project.get("title", "Untitled Project"), 10, bold=True)
            for bullet in project.get("bullets", [])[:4]:
                if str(bullet).strip():
                    _bullet_to_cell(right, bullet, 9)
            if project.get("tech_stack"):
                _cell_text(right, f"Stack: {project['tech_stack']}", 8, italic=True)
    else:
        for exp in experience[:4]:
            _cell_text(right, f"{exp.get('role', '')}, {exp.get('company', '')}".strip(", "), 10, bold=True)
            for bullet in exp.get("bullets", [])[:4]:
                if str(bullet).strip():
                    _bullet_to_cell(right, bullet, 9)
        if projects:
            _cell_heading(right, "Projects", accent)
            for project in projects[:3]:
                _cell_text(right, project.get("title", "Untitled Project"), 9, bold=True)
                for bullet in project.get("bullets", [])[:2]:
                    if str(bullet).strip():
                        _bullet_to_cell(right, bullet, 8)

    if data.get("certifications"):
        _cell_heading(right, "Certifications", accent)
        for cert in data["certifications"][:5]:
            _cell_text(right, cert.get("name", ""), 8)
    return doc


def _build_campus(data: Dict[str, Any]) -> Document:
    doc = Document()
    _margins(doc)
    personal = data.get("personal", {})
    p = _para(doc, sp_before=0, sp_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    _run(p, personal.get("name") or "[Your Full Name]", sz=Pt(18), bold=True)
    _run(_para(doc, sp_after=5), " | ".join([x for x in [personal.get("email"), personal.get("phone"), personal.get("location")] if x]), sz=Pt(9))
    if data.get("summary"):
        _heading(doc, "Campus Recruiter Summary")
        _run(_para(doc, sp_after=3), data["summary"], sz=Pt(10))
    _heading(doc, "Education")
    current = personal.get("current_education") or {}
    if current.get("institution"):
        _run(_para(doc, sp_after=2), f"{current.get('degree', 'B.Tech')} {current.get('branch', '')}, {current.get('institution', '')}".strip(", "), sz=Pt(11), bold=True)
    _heading(doc, "Placement Projects")
    for project in data.get("projects", [])[:4]:
        _run(_para(doc, sp_after=1), project.get("title", "Untitled Project"), sz=Pt(11), bold=True)
        for bullet in project.get("bullets", [])[:3]:
            if str(bullet).strip():
                _bullet(doc, bullet)
    _heading(doc, "Technical Skills")
    skills = data.get("skills", {})
    for key, label in [("languages", "Languages"), ("frameworks", "Frameworks"), ("tools", "Tools"), ("databases", "Databases")]:
        if skills.get(key):
            p = _para(doc, sp_after=1)
            _run(p, f"{label}: ", sz=Pt(10), bold=True)
            _run(p, ", ".join(skills[key]), sz=Pt(10))
    if data.get("experience"):
        _heading(doc, "Experience")
        for exp in data["experience"][:3]:
            _run(_para(doc, sp_after=1), f"{exp.get('role', '')}, {exp.get('company', '')}".strip(", "), sz=Pt(10), bold=True)
    return doc


def _build_by_template(data: Dict[str, Any], template: str) -> Document:
    if template == "compact":
        return _build_compact(data)
    if template == "campus":
        return _build_campus(data)
    if template in {"modern", "developer", "core-engineering", "management"}:
        return _build_sidebar(data, template)
    return _build_classic(data)


# ═══════════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════════

async def generate_docx(
    user: dict,
    session: AsyncSession,
    template: str = "classic",
) -> io.BytesIO:
    """Generate a resume .docx from the student's resume profile data."""
    result = await session.execute(
        select(models.UserProfile).where(
            models.UserProfile.user_id == user["id"],
        )
    )
    profile = result.scalars().first()
    if not profile:
        raise ValueError("No profile found")

    extra = profile.extra_data or {}
    resume = extra.get("resume_profile", {})

    # Resolve college name
    user_obj = await session.get(models.User, user["id"])
    institution = ""
    if user_obj and user_obj.college_id:
        college = await session.get(models.College, user_obj.college_id)
        if college:
            institution = college.name

    # Get user's name from User model (not UserProfile)
    student_name_str = user_obj.name if user_obj else user.get("full_name", "Resume")

    data = {
        "personal": {
            "name": student_name_str,
            "email": resume.get("email") or user.get("email", ""),
            "phone": resume.get("phone") or getattr(profile, "phone", "") or "",
            "location": resume.get("location", ""),
            "linkedin": resume.get("linkedin", ""),
            "github": resume.get("github", ""),
            "portfolio": resume.get("portfolio", ""),
            "current_education": {
                "degree": "B.Tech",
                "institution": institution,
                "branch": profile.department or "",
                "batch": profile.batch or "",
            },
        },
        "summary": resume.get("summary", ""),
        "education_history": resume.get("education_history", []),
        "skills": resume.get("skills", {}),
        "projects": resume.get("projects", []),
        "experience": resume.get("experience", []),
        "certifications": resume.get("certifications", []),
        "achievements": resume.get("achievements", []),
    }

    doc = _build_by_template(data, template)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = (student_name_str or "resume").replace(" ", "_")
    logger.info("Generated %s resume for student %s", template, user["id"])
    return buffer, f"{safe_name}_Resume.docx"
